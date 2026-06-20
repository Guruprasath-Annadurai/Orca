"""
Text extraction from uploaded documents.

Supported formats:
  .pdf   → pdfminer.six (pure Python, no poppler) or PyPDF2 fallback
  .txt / .md / .rst → direct UTF-8 read
  .docx  → python-docx
  .csv   → csv module → markdown table
  .json  → pretty-printed subset

All extraction is local — no external APIs.
"""
from __future__ import annotations

import csv
import io
import json
import re
from pathlib import Path


def extract(filename: str, data: bytes) -> str:
    """Extract plain text from file bytes. Returns clean text."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(data)
    elif ext in (".txt", ".md", ".rst", ".log", ".yaml", ".yml", ".toml", ".ini", ".env"):
        return data.decode("utf-8", errors="replace").strip()
    elif ext == ".docx":
        return _extract_docx(data)
    elif ext == ".csv":
        return _extract_csv(data)
    elif ext == ".json":
        return _extract_json(data)
    elif ext in (".py", ".js", ".ts", ".java", ".cpp", ".c", ".go", ".rb", ".rs", ".sh"):
        lang = ext.lstrip(".")
        text = data.decode("utf-8", errors="replace").strip()
        return f"```{lang}\n{text}\n```"
    else:
        # Try UTF-8 decode as last resort
        try:
            return data.decode("utf-8", errors="replace").strip()
        except Exception:
            raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(data: bytes) -> str:
    # Try pdfminer first (better quality)
    try:
        from pdfminer.high_level import extract_text_to_fp
        from pdfminer.layout import LAParams
        out = io.StringIO()
        extract_text_to_fp(io.BytesIO(data), out, laparams=LAParams())
        text = out.getvalue().strip()
        if text:
            return _clean_pdf_text(text)
    except ImportError:
        pass

    # Fall back to PyPDF2
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        text = "\n\n".join(pages).strip()
        if text:
            return _clean_pdf_text(text)
    except ImportError:
        pass

    # Last resort: pypdf (newer fork of PyPDF2)
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n\n".join(pages).strip()
        return _clean_pdf_text(text)
    except ImportError:
        raise ImportError(
            "PDF extraction requires: pip install pypdf\n"
            "Or for better quality: pip install pdfminer.six"
        )


def _clean_pdf_text(text: str) -> str:
    """Remove common PDF extraction artifacts."""
    text = re.sub(r"\n{3,}", "\n\n", text)     # collapse excess blank lines
    text = re.sub(r"[ \t]{2,}", " ", text)      # collapse horizontal whitespace
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text) # join hyphenated line breaks
    return text.strip()


def _extract_docx(data: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab table cells
        for table in doc.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n\n".join(paragraphs)
    except ImportError:
        raise ImportError("DOCX extraction requires: pip install python-docx")


def _extract_csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    reader = csv.DictReader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return text

    headers = list(rows[0].keys())
    lines = ["| " + " | ".join(headers) + " |"]
    lines.append("| " + " | ".join("---" for _ in headers) + " |")
    for row in rows[:200]:  # cap at 200 rows for token budget
        lines.append("| " + " | ".join(str(row.get(h, "")) for h in headers) + " |")
    if len(rows) > 200:
        lines.append(f"\n_(showing 200 of {len(rows)} rows)_")
    return "\n".join(lines)


def _extract_json(data: bytes) -> str:
    try:
        obj = json.loads(data)
        # Pretty-print but cap size
        pretty = json.dumps(obj, indent=2)
        if len(pretty) > 8000:
            pretty = pretty[:8000] + "\n...(truncated)"
        return f"```json\n{pretty}\n```"
    except Exception:
        return data.decode("utf-8", errors="replace")


SUPPORTED_EXTENSIONS = {
    ".pdf", ".txt", ".md", ".rst", ".log",
    ".yaml", ".yml", ".toml", ".ini",
    ".docx", ".csv", ".json",
    ".py", ".js", ".ts", ".java", ".cpp", ".c",
    ".go", ".rb", ".rs", ".sh",
}

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
